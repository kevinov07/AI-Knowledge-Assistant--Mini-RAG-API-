"""
Procesamiento de documentos por extensión y normalización de texto.
Incluye limpieza específica para PDF (menos fragmentación, mejor calidad para RAG/embeddings).
"""
import re
import io

import pandas as pd
import pypdf
from docx import Document as DocxDocument

import fitz

def normalize_text_for_rag(text: str) -> str:
    """
    Normaliza texto extraído (sobre todo de PDF) para mejorar calidad en chunks y búsqueda.
    - Une líneas fragmentadas típicas de PDF (cada palabra en una línea).
    - Colapsa espacios y saltos excesivos.
    - Preserva párrafos (doble salto) y mejora coherencia para embeddings.
    """
    if not text or not text.strip():
        return ""
    # Normalizar saltos de línea a \n
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    if not lines:
        return ""

    # Detectar líneas que son continuación (no terminan en . ! ? " y la siguiente no empieza mayúscula)
    sentence_end = re.compile(r"[.!?]\s*$")
    merged = []
    i = 0
    while i < len(lines):
        current = lines[i]
        # Unir líneas fragmentadas: si la línea es corta y la siguiente no inicia párrafo, unir
        while i + 1 < len(lines):
            next_ln = lines[i + 1]
            current_ends_sentence = bool(sentence_end.search(current))
            next_starts_upper = len(next_ln) > 0 and next_ln[0].isupper()
            # Si la actual no termina en punto y la siguiente es corta o no empieza mayúscula, unir
            is_short = len(next_ln.split()) <= 4 or len(next_ln) < 40
            if not current_ends_sentence and (not next_starts_upper or is_short):
                current = f"{current} {next_ln}".strip()
                i += 1
            elif current_ends_sentence and next_starts_upper:
                break
            else:
                current = f"{current} {next_ln}".strip()
                i += 1
        merged.append(current)
        i += 1

    # Un párrafo = bloque de líneas; separar por doble salto donde haya cambio claro de párrafo
    result = []
    for j, line in enumerate(merged):
        result.append(line)
        # Añadir salto de párrafo si la siguiente empieza por mayúscula y la actual termina en punto
        if j + 1 < len(merged) and sentence_end.search(line) and merged[j + 1] and merged[j + 1][0].isupper():
            result.append("")
    text = "\n".join(result)
    # Colapsar múltiples espacios y múltiples saltos
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class DocumentProcessor:
    """
    Clase para procesar documentos de diferentes formatos y extraer texto.
    Soporta: .txt, .pdf, .docx, .md, .csv, .xlsx, .xls
    """
    
    SUPPORTED_FORMATS = {'.txt', '.pdf', '.docx', '.md', '.csv', '.xlsx', '.xls'}
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Verifica si el formato del archivo es soportado."""
        extension = filename.lower().split('.')[-1] if '.' in filename else ''
        return f'.{extension}' in DocumentProcessor.SUPPORTED_FORMATS
    
    @staticmethod
    def get_extension(filename: str) -> str:
        """Obtiene la extensión del archivo."""
        return f'.{filename.lower().split(".")[-1]}' if '.' in filename else ''
    
    @staticmethod
    def process_txt(content: bytes) -> str:
        """Procesa archivos de texto plano (.txt)."""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError as e:
            # Intenta con otras codificaciones comunes
            last_error = e
            for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError as e2:
                    last_error = e2
                    continue
            raise ValueError("No se pudo decodificar el archivo con ninguna codificación conocida") from last_error
    
    @staticmethod
    def _extract_text_pymupdf(content: bytes) -> str:
        """Extrae texto con PyMuPDF (principal para PDFs)."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            parts = [page.get_text().strip() for page in doc if page.get_text().strip()]
            doc.close()
            return "\n".join(parts)
        except Exception:
            return ""

    @staticmethod
    def _extract_text_pypdf(content: bytes) -> str:
        """Extrae texto con pypdf (pure Python; mismo motor que PyPDF2)."""
        try:
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file, strict=False)
            if getattr(reader, "is_encrypted", False) and reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    pass
            parts = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        parts.append(text.strip())
                except Exception:
                    continue
            return "\n".join(parts)
        except Exception:
            return ""

    @staticmethod
    def process_pdf(content: bytes) -> str:
        """
        Procesa PDF: PyMuPDF como principal, pypdf como fallback.
        """
        text = DocumentProcessor._extract_text_pymupdf(content)
        if text.strip():
            return text
        text = DocumentProcessor._extract_text_pypdf(content)
        if text.strip():
            return text
        raise ValueError(
            "No se extrajo texto del PDF: las páginas parecen ser solo imágenes (PDF escaneado). "
            "Sube un PDF con texto seleccionable o convierte el documento a texto antes de subirlo."
        )

    @staticmethod
    def process_docx(content: bytes) -> str:
        """Procesa archivos Word (.docx)."""
        try:
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            
            text_parts = []
            
            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("El documento DOCX está vacío")
            
            return full_text
        except Exception as e:
            raise ValueError(f"Error al procesar DOCX: {str(e)}") from e
    
    @staticmethod
    def process_markdown(content: bytes) -> str:
        """Procesa archivos Markdown (.md)."""
        # Para Markdown, podemos simplemente leerlo como texto plano
        # ya que el formato Markdown es legible en texto plano
        # Si quieres convertirlo a HTML primero, puedes usar la librería markdown
        try:
            text = content.decode('utf-8')
            
            # Opcionalmente, puedes usar la librería markdown para procesarlo
            # import markdown
            # html = markdown.markdown(text)
            # # Luego convertir HTML a texto plano con html2text o similar
            
            if not text.strip():
                raise ValueError("El archivo Markdown está vacío")
            
            return text
        except UnicodeDecodeError as e:
            raise ValueError("No se pudo decodificar el archivo Markdown") from e
    
    @staticmethod
    def process_csv(content: bytes) -> str:
        """Procesa archivos CSV (.csv)."""
        try:
            csv_file = io.BytesIO(content)
            
            # Intentar leer el CSV con diferentes delimitadores
            df = None
            for delimiter in [',', ';', '\t', '|']:
                try:
                    csv_file.seek(0)
                    df = pd.read_csv(csv_file, delimiter=delimiter, encoding='utf-8')
                    if len(df.columns) > 1:  # Si tiene más de una columna, probablemente es el delimitador correcto
                        break
                except:
                    continue
            
            if df is None or df.empty:
                # Intentar con otras codificaciones
                for encoding in ['latin-1', 'iso-8859-1', 'cp1252']:
                    try:
                        csv_file.seek(0)
                        df = pd.read_csv(csv_file, encoding=encoding)
                        break
                    except:
                        continue
            
            if df is None or df.empty:
                raise ValueError("No se pudo leer el archivo CSV o está vacío")
            
            # Convertir el DataFrame a texto de manera estructurada
            text_parts = []
            
            # Agregar encabezados
            headers = ' | '.join(str(col) for col in df.columns)
            text_parts.append(f"Columnas: {headers}\n")
            
            # Agregar filas
            for idx, row in df.iterrows():
                row_text = ' | '.join(f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val))
                if row_text:
                    text_parts.append(f"Fila {idx + 1}: {row_text}")
            
            full_text = '\n'.join(text_parts)
            
            if not full_text.strip():
                raise ValueError("El archivo CSV no contiene datos")
            
            return full_text
        except Exception as e:
            raise ValueError(f"Error al procesar CSV: {str(e)}") from e
    
    @staticmethod
    def process_excel(content: bytes) -> str:
        """Procesa archivos Excel (.xlsx, .xls)."""
        try:
            excel_file = io.BytesIO(content)
            try:
                df = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            except Exception:
                excel_file.seek(0)
                df = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
            if not df or all(d.empty for d in df.values()):
                raise ValueError("El archivo Excel está vacío o no se pudo leer")
            text_parts = []
            for sheet_name, sheet_df in df.items():
                if sheet_df.empty:
                    continue
                text_parts.append(f"\n=== Hoja: {sheet_name} ===\n")
                headers = ' | '.join(str(col) for col in sheet_df.columns)
                text_parts.append(f"Columnas: {headers}\n")
                for idx, row in sheet_df.iterrows():
                    row_text = ' | '.join(f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val))
                    if row_text:
                        text_parts.append(f"Fila {idx + 1}: {row_text}")
            full_text = '\n'.join(text_parts)
            if not full_text.strip():
                raise ValueError("El archivo Excel no contiene datos")
            return full_text
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Error al procesar Excel: {str(e)}") from e
    
    @staticmethod
    def process_document(content: bytes, filename: str) -> str:
        """
        Procesa un documento y extrae su texto según su formato.
        
        Args:
            content: Contenido del archivo en bytes
            filename: Nombre del archivo con extensión
            
        Returns:
            str: Texto extraído del documento
            
        Raises:
            ValueError: Si el formato no es soportado o hay error al procesar
        """
        extension = DocumentProcessor.get_extension(filename)
        
        if not DocumentProcessor.is_supported(filename):
            raise ValueError(
                f"Formato {extension} no soportado. "
                f"Formatos soportados: {', '.join(sorted(DocumentProcessor.SUPPORTED_FORMATS))}"
            )
        
        processors = {
            '.txt': DocumentProcessor.process_txt,
            '.pdf': DocumentProcessor.process_pdf,
            '.docx': DocumentProcessor.process_docx,
            '.md': DocumentProcessor.process_markdown,
            '.csv': DocumentProcessor.process_csv,
            '.xlsx': DocumentProcessor.process_excel,
            '.xls': DocumentProcessor.process_excel,
        }
        
        processor = processors.get(extension)
        if not processor:
            raise ValueError(f"No hay procesador disponible para {extension}")

        raw = processor(content)
        # Aplicar normalización a todos los formatos para consistencia y mejor RAG
        return normalize_text_for_rag(raw)
